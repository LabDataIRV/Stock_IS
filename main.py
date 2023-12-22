import sys
import os
import sqlite3
from PySide2 import QtCore, QtWidgets
from ui_stock import Ui_MainWindow
from ui_delivery import Ui_ProvideWindow
from ui_ship import Ui_ShopWindow
from ui_return import Ui_ReturnWindow
from docx import Document


# Main stock window
class window(QtWidgets.QMainWindow):
    def __init__(self):
        super(window, self).__init__()
        # GUI connections
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.deliveryB.clicked.connect(self.onDelivery)
        self.ui.shipB.clicked.connect(self.onShip)
        self.ui.returnB.clicked.connect(self.onReturn)
        self.ui.deliveryDocB.clicked.connect(self.onDeliveryDoc)
        self.ui.shipDocB.clicked.connect(self.onShipmentDoc)
        self.ui.returnDocB.clicked.connect(self.onReturnDoc)
        self.updateState()


    def updateBatches(self):
        connection, cursor = self.connectDB()
        cursor.execute("""
                SELECT DISTINCT batch 
                FROM stock 
                """)
        deliv_batches = cursor.fetchall()
        cursor.execute("""
                        SELECT DISTINCT batch 
                        FROM operations_shops
                        """)
        ship_batches = cursor.fetchall()
        connection.close()
        self.ui.deliveryBatchCB.clear()
        print(str(deliv_batches[0][0]))
        for batch in deliv_batches:
            self.ui.deliveryBatchCB.addItem(str(batch[0]))
        self.ui.shipBatchCB.clear()
        for batch in ship_batches:
            self.ui.shipBatchCB.addItem(str(batch[0]))

    # delivery batch Document button have been pressed
    def onDeliveryDoc(self):
        batch_name = self.ui.deliveryBatchCB.currentText()

        connection, cursor = self.connectDB()
        cursor.execute("""
                        SELECT material.name_material, quantity, price, provider.name_provider
                        FROM stock JOIN provider ON stock.provider = provider.id
                        JOIN material ON stock.material = material.id
                        WHERE batch = '""" + batch_name + "'"
                        )
        deliv_batch = cursor.fetchall()
        connection.close()
        delivDoc = []
        sumQ = 0
        sumP = 0
        for i, pos in enumerate(deliv_batch):
            delivDoc.append([i+1, pos[0], pos[3], pos[1], pos[2]]) # N material provider quantity price
            sumQ += float(pos[1])
            sumP += float(pos[2])

        document = Document()
        document.add_heading('Накладная поставки товара', 0)
        p = document.add_paragraph('№ ')
        p.add_run(batch_name).bold = True
        document.add_heading('Состав партии', level=1)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Материал'
        hdr_cells[2].text = 'Поставщик'
        hdr_cells[3].text = 'Количество'
        hdr_cells[4].text = 'Стоимость'
        for n, mat, prov, Q, P in delivDoc:
            row_cells = table.add_row().cells
            row_cells[0].text = str(n)
            row_cells[1].text = mat
            row_cells[2].text = prov
            row_cells[3].text = str(Q)
            row_cells[4].text = str(P)
        p = document.add_paragraph('Поставлено '+ str(sumQ) + ' единиц товара на сумму '+str(sumP))
        document.add_page_break()
        document.save('DeliveryReport_'+batch_name+'.docx')
        self.ui.statusbar.showMessage("Накладная готова.")
        os.system('start DeliveryReport_'+batch_name+'.docx')

    # shipment batch document button have been pressed
    def onShipmentDoc(self):
        batch_name = self.ui.shipBatchCB.currentText()

        connection, cursor = self.connectDB()
        cursor.execute("""
                               SELECT material.name_material, shops.number_shops, operations_shops.quantity, operations_shops.price
                               FROM operations_shops JOIN material ON operations_shops.material = material.id
                               JOIN shops ON operations_shops.number_shops = shops.id
                               WHERE operation = 1 AND batch = '""" + batch_name + "'"
                       )
        ship_batch = cursor.fetchall()
        connection.close()
        shipDoc = []
        sumQ = 0
        sumP = 0
        for i, pos in enumerate(ship_batch):
            shipDoc.append([i + 1, pos[0], pos[1], pos[2], pos[3]])  # N material shop quantity price
            sumQ += float(pos[2])
            sumP += float(pos[3])

        document = Document()
        document.add_heading('Накладная отгрузки товара', 0)
        p = document.add_paragraph('№ ')
        p.add_run(batch_name).bold = True
        document.add_heading('Состав партии', level=1)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Материал'
        hdr_cells[2].text = 'Магазин'
        hdr_cells[3].text = 'Количество'
        hdr_cells[4].text = 'Стоимость'
        for n, mat, shop, Q, P in shipDoc:
            row_cells = table.add_row().cells
            row_cells[0].text = str(n)
            row_cells[1].text = mat
            row_cells[2].text = shop
            row_cells[3].text = str(Q)
            row_cells[4].text = str(P)
        p = document.add_paragraph('Отгружено ' + str(sumQ) + ' единиц товара на сумму ' + str(sumP))
        document.add_page_break()
        document.save('ShipmentReport_' + batch_name + '.docx')
        self.ui.statusbar.showMessage("Накладная готова.")

    # return batch document button have been pressed
    def onReturnDoc(self):
        batch_name = self.ui.shipBatchCB.currentText()

        connection, cursor = self.connectDB()
        cursor.execute("""
                                       SELECT material.name_material, shops.number_shops, operations_shops.quantity, operations_shops.price
                                       FROM operations_shops JOIN material ON operations_shops.material = material.id
                                       JOIN shops ON operations_shops.number_shops = shops.id
                                       WHERE operation = -1 AND batch = '""" + batch_name + "'"
                       )
        ship_batch = cursor.fetchall()
        connection.close()
        print(len(ship_batch))
        if len(ship_batch) <=0:
            self.ui.statusbar.showMessage("Возвратов не было!.")
            return
        shipDoc = []
        sumQ = 0
        sumP = 0
        for i, pos in enumerate(ship_batch):
            shipDoc.append([i + 1, pos[0], pos[1], pos[2], pos[3]])  # N material shop quantity price
            sumQ += float(pos[2])
            sumP += float(pos[3])

        document = Document()
        document.add_heading('Накладная возврата товара', 0)
        p = document.add_paragraph('№ ')
        p.add_run(batch_name).bold = True
        document.add_heading('Возвращенный товар', level=1)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Материал'
        hdr_cells[2].text = 'Магазин'
        hdr_cells[3].text = 'Количество'
        hdr_cells[4].text = 'Стоимость'
        for n, mat, shop, Q, P in shipDoc:
            row_cells = table.add_row().cells
            row_cells[0].text = str(n)
            row_cells[1].text = mat
            row_cells[2].text = shop
            row_cells[3].text = str(Q)
            row_cells[4].text = str(P)
        p = document.add_paragraph('Возвращено ' + str(sumQ) + ' единиц товара на сумму ' + str(sumP))
        document.add_page_break()
        document.save('ReturnsReport_' + batch_name + '.docx')
        self.ui.statusbar.showMessage("Накладная готова.")

    # Belivery button have been pressed
    def onDelivery(self):
        self.deliver = delive_win()
        self.deliver.show()
        application.hide()
    # Shipment button have been pressed
    def onShip(self):
        self.ship = ship_win()
        self.ship.show()
        application.hide()
    # Return button have been pressed
    def onReturn(self):
        self.ret = return_win()
        self.ret.show()
        application.hide()
    # Connect to database
    def connectDB(self):
        connection = sqlite3.connect('stock.db')
        cursor = connection.cursor()
        return connection, cursor
    # Update stock table
    def updateState(self):
        # Table setting
        self.ui.stockT.clear()
        self.ui.stockT.setRowCount(1)
        self.ui.stockT.setColumnCount(3)
        self.ui.stockT.setItem(0, 0, QtWidgets.QTableWidgetItem("Материал"))
        self.ui.stockT.setItem(0, 1, QtWidgets.QTableWidgetItem("Количество"))
        self.ui.stockT.setItem(0, 2, QtWidgets.QTableWidgetItem("Средняя цена за ед."))
        # Take list of materials
        materials = [] # name, Qdelvery, avg_price
        connection, cursor = self.connectDB()
        cursor.execute("""
        SELECT material.id, material.name_material, SUM(stock.quantity), AVG(stock.price) 
        FROM stock JOIN material ON material.id = stock.material 
        GROUP BY material.name_material
        """)
        names = cursor.fetchall()
        for name in names:
            materials.append([name[1], name[2], name[3]])
            # Take information about shipment and returns
            cursor.execute("SELECT quantity, operation FROM operations_shops WHERE material = "+str(name[0]))
            operations = cursor.fetchall()
            for ops in operations:
                materials[-1][1] -= ops[0]*ops[1] # Q Corrections
            if materials[-1][1] > 0: materials[-1][2] /= materials[-1][1]
            else: materials[-1][2] = 0
        connection.close()
        self.ui.stockT.setRowCount(len(materials)+1)
        for row, name in enumerate(materials):
            self.ui.stockT.setItem(row + 1, 0, QtWidgets.QTableWidgetItem(str(name[0])))
            self.ui.stockT.setItem(row + 1, 1, QtWidgets.QTableWidgetItem(str(name[1])))
            self.ui.stockT.setItem(row + 1, 2, QtWidgets.QTableWidgetItem("%.2f" % name[2]))
        self.ui.stockT.resizeColumnsToContents()
        self.updateBatches()

# Delivery window
class delive_win(QtWidgets.QMainWindow):
    def __init__(self):
        super(delive_win, self).__init__()
        self.uiDeliver = Ui_ProvideWindow()
        self.uiDeliver.setupUi(self)
        self.uiDeliver.deliveryB.clicked.connect(self.onDelivery)
        self.uiDeliver.addProviderB.clicked.connect(self.onAddProvider)
        self.uiDeliver.addMaterialB.clicked.connect(self.onAddMaterial)
        self.updateProvider()
        self.updateMaterial()
    # Close delivery window
    def closeEvent(self, event):
        application.show()
    # Connect to database
    def connectDB(self):
        connection = sqlite3.connect('stock.db')
        cursor = connection.cursor()
        return connection, cursor
    # Add new Provider button have been pressed
    def onAddProvider(self):
        newProvider = self.uiDeliver.providerE.text()
        # проверка ввода не пустого имени поставщика
        if newProvider == '':
            self.uiDeliver.statusbar.showMessage("Введите имя поставщика!")
            return
        connection, cursor = self.connectDB()
        # проверка уникальности введенного имени
        cursor.execute("SELECT * FROM provider WHERE name_provider = '"+newProvider+"'")
        names = cursor.fetchall()
        if len(names)==0:
            #запрос
            cursor.execute("INSERT INTO provider (name_provider) VALUES ('"+newProvider+"')")
            connection.commit()
            self.uiDeliver.statusbar.showMessage("Поставщик добавлен.")
        else: # ошибка
            self.uiDeliver.statusbar.showMessage("Поставщик уже был добавлен!")
        connection.close()
        self.updateProvider()
    # Update list of provides
    def updateProvider(self):
        connection, cursor = self.connectDB()
        cursor.execute("SELECT name_provider FROM provider")
        names = cursor.fetchall()
        connection.close()
        self.uiDeliver.providerCB.clear()
        for name in names:
            self.uiDeliver.providerCB.addItem(str(name[0]))
    # Add new material button have been pressed
    def onAddMaterial(self):
        newMaterial = self.uiDeliver.materialE.text()
        # проверка ввода не пустого имени поставщика
        if newMaterial == '':
            self.uiDeliver.statusbar.showMessage("Введите название материала!")
            return
        connection, cursor = self.connectDB()
        # проверка уникальности введенного имени
        cursor.execute("SELECT * FROM material WHERE name_material = '"+newMaterial+"'")
        names = cursor.fetchall()
        if len(names)==0:
            #запрос
            cursor.execute("INSERT INTO material (name_material) VALUES ('"+newMaterial+"')")
            connection.commit()
            self.uiDeliver.statusbar.showMessage("Материал добавлен.")
        else: # ошибка
            self.uiDeliver.statusbar.showMessage("Материал уже был добавлен!")
        connection.close()
        self.updateMaterial()
    # Update lis of materials
    def updateMaterial(self):
        connection, cursor = self.connectDB()
        cursor.execute("SELECT name_material FROM material")
        names = cursor.fetchall()
        connection.close()
        self.uiDeliver.materialCB.clear()
        for name in names:
            self.uiDeliver.materialCB.addItem(str(name[0]))
    # Delivery button have been pressed
    def onDelivery(self):
        material = self.uiDeliver.materialCB.currentText()
        provider = self.uiDeliver.providerCB.currentText()
        quantity = self.uiDeliver.quantityE.text()
        price = self.uiDeliver.priceE.text()
        batch = self.uiDeliver.batchE.text()
        # Проверка корректности ввода количества, цены и номера партии
        try:
            if int(quantity)<=0:
                self.uiDeliver.statusbar.showMessage("Количество должно быть больше 0")
                return
            if float(price)<0:
                self.uiDeliver.statusbar.showMessage("Цена должна быть положительная")
                return
            if batch =='':
                self.uiDeliver.statusbar.showMessage("Введите номер партии")
                return
        except:
            self.uiDeliver.statusbar.showMessage("Ошибка во введенных данных!")
            return
        connection, cursor = self.connectDB()
        # проверка существования поставщика
        cursor.execute("SELECT * FROM provider WHERE name_provider = '" + provider + "'")
        names = cursor.fetchall()
        print(names)
        provider_id = -1
        if len(names) == 1:
            provider_id = names[0][0]
            print(provider_id)
        # проверка существования материала
        cursor.execute("SELECT * FROM material WHERE name_material = '" + material + "'")
        names = cursor.fetchall()
        print(names)
        material_id = -1
        if len(names) == 1:
            material_id = names[0][0]
            print(material_id)
        # проверка существования поставки материала в данной партии
        cursor.execute("SELECT * FROM stock WHERE material = '" + str(material_id) + "' AND batch = '"+ batch+"'")
        names = cursor.fetchall()
        print(names)
        if len(names) > 0:
            self.uiDeliver.statusbar.showMessage("В данной партии этот материал уже введен!")
            return
        connection.close()
        # занесение сведений в БД
        connection, cursor = self.connectDB()
        cursor.execute("INSERT INTO stock (material, quantity, price, batch, provider) VALUES (" + str(material_id) +
                                                                                               ","+quantity +
                                                                                               ","+price +
                                                                                               ",'" + batch +
                                                                                               "'," + str(provider_id) +
                                                                                              ")")
        connection.commit()
        connection.close()
        application.deliver.hide()
        application.updateState()
        application.show()

# Shipment window
class ship_win(QtWidgets.QMainWindow):
    def __init__(self):
        super(ship_win, self).__init__()
        self.uiShip = Ui_ShopWindow()
        self.uiShip.setupUi(self)
        self.uiShip.shipB.clicked.connect(self.onShiping)
        self.uiShip.addShopB.clicked.connect(self.onAddShop)
        self.updateShop()
        self.updateMaterial()
    # Close shipment window
    def closeEvent(self, event):
        application.show()
    # Shipment button have been pressed
    def onShiping(self):
        material = self.uiShip.materialCB.currentText()
        shop = self.uiShip.shopsCB.currentText()
        quantity = self.uiShip.quantityE.text()
        price = self.uiShip.priceE.text()
        batch = self.uiShip.batchE.text()
        # Проверка корректности ввода количества, цены и номера партии
        try:
            if int(quantity) <= 0:
                self.uiShip.statusbar.showMessage("Количество должно быть больше 0")
                return
            if float(price) < 0:
                self.uiShip.statusbar.showMessage("Цена должна быть положительная")
                return
            if batch == '':
                self.uiShip.statusbar.showMessage("Введите номер партии")
                return
        except:
            self.uiShip.statusbar.showMessage("Ошибка во введенных данных!")
            return
        connection, cursor = self.connectDB()
        # проверка существования магазина
        cursor.execute("SELECT * FROM shops WHERE number_shops = '" + shop + "'")
        names = cursor.fetchall()
        print(names)
        shop_id = -1
        if len(names) == 1:
            shop_id = names[0][0]
            print(shop_id)
        # проверка существования материала
        cursor.execute("SELECT * FROM material WHERE name_material = '" + material + "'")
        names = cursor.fetchall()
        print(names)
        material_id = -1
        if len(names) == 1:
            material_id = names[0][0]
            print(material_id)
        # проверка наличия материала
        cursor.execute("""
               SELECT SUM(stock.quantity) 
               FROM stock JOIN material ON material.id = stock.material
               WHERE  material.id = """ + str(material_id) +
               """ GROUP BY material.name_material
               """)
        names = cursor.fetchall()
        max_quantity = 0
        if len(names) == 1:
            max_quantity = names[0][0]
            # Take information about shipment and returns
            cursor.execute("SELECT quantity, operation FROM operations_shops WHERE material = " + str(material_id))
            operations = cursor.fetchall()
            for ops in operations:
                max_quantity -= ops[0] * ops[1]  # Q Corrections
        # проверка существования поставки материала в данной партии
        cursor.execute(
                "SELECT * FROM operations_shops WHERE material = '" + str(material_id) + "' AND batch = '" + batch + "' AND operation = 1")
        names = cursor.fetchall()
        connection.close()
        print(names)
        if len(names) > 0:
               self.uiShip.statusbar.showMessage("В данной партии этот материал уже введен!")
               return
        if int(quantity) > max_quantity:
            self.uiShip.statusbar.showMessage(
                "Можно поставить не более " + str(max_quantity))
            return

        # занесение сведений в БД
        connection, cursor = self.connectDB()
        cursor.execute("INSERT INTO operations_shops (number_shops, material, quantity, price, batch, operation) VALUES (" +
                           str(shop_id) +
                           "," + str(material_id) +
                           "," + quantity +
                           "," + price +
                           ",'" + batch +
                           "',1)")
        connection.commit()
        connection.close()
        application.ship.hide()
        application.updateState()
        application.show()
    # Add new shop button have been pressed
    def onAddShop(self):
        newShop = self.uiShip.shopE.text()
        # проверка ввода не пустого имени поставщика
        if newShop == '':
            self.uiShip.statusbar.showMessage("Введите имя магазина!")
            return
        connection, cursor = self.connectDB()
        # проверка уникальности введенного имени
        cursor.execute("SELECT * FROM shops WHERE number_shops = '" + newShop + "'")
        names = cursor.fetchall()
        if len(names) == 0:
            # запрос
            cursor.execute("INSERT INTO shops (number_shops) VALUES ('" + newShop + "')")
            connection.commit()
            self.uiShip.statusbar.showMessage("Магазин добавлен.")
        else:  # ошибка
            self.uiShip.statusbar.showMessage("Магазин уже был добавлен!")
        connection.close()
        self.updateShop()
    # Connect to database
    def connectDB(self):
        connection = sqlite3.connect('stock.db')
        cursor = connection.cursor()
        return connection, cursor
    # Update list of shops
    def updateShop(self):
        connection, cursor = self.connectDB()
        cursor.execute("SELECT number_shops FROM shops")
        names = cursor.fetchall()
        connection.close()
        self.uiShip.shopsCB.clear()
        for name in names:
            self.uiShip.shopsCB.addItem(str(name[0]))
    # Uptate list of materials
    def updateMaterial(self):
        connection, cursor = self.connectDB()
        cursor.execute("SELECT name_material FROM material")
        names = cursor.fetchall()
        connection.close()
        self.uiShip.materialCB.clear()
        for name in names:
            self.uiShip.materialCB.addItem(str(name[0]))

# Return window
class return_win(QtWidgets.QMainWindow):
    def __init__(self):
        super(return_win, self).__init__()
        self.uiReturn = Ui_ReturnWindow()
        self.uiReturn.setupUi(self)
        self.uiReturn.returnB.clicked.connect(self.onReturn)
        self.uiReturn.shopsCB.currentIndexChanged.connect(self.onShopChange)
        self.updateShop()
        if self.uiReturn.shopsCB.count()>0:
            self.updateBatch(self.uiReturn.shopsCB.currentText())
        else:
            self.updateBatch()

    def closeEvent(self, event):
        application.show()

    def connectDB(self):
        connection = sqlite3.connect('stock.db')
        cursor = connection.cursor()
        return connection, cursor

    def onShopChange(self, index):
        self.updateBatch(self.uiReturn.shopsCB.currentText())

    def onReturn(self): #1 Переделать! 2 Нельзя принять больше размера партии и прошлых возвратов!
        if self.uiReturn.batchT.currentRow()<1:
            self.uiReturn.statusbar.showMessage("Поставка не выбрана!")
            return
        batch    = self.uiReturn.batchT.item(self.uiReturn.batchT.currentRow(), 0).text()
        material = self.uiReturn.batchT.item(self.uiReturn.batchT.currentRow(), 1).text()
        shop_name = self.uiReturn.shopsCB.currentText()
        max_quantity = self.uiReturn.batchT.item(self.uiReturn.batchT.currentRow(), 2).text()
        rets_quantity = self.uiReturn.batchT.item(self.uiReturn.batchT.currentRow(), 3).text()
        quantity = self.uiReturn.quantityE.text()
        # Проверка корректности ввода
        try:
            if int(quantity) <= 0:
                self.uiReturn.statusbar.showMessage("Количество должно быть больше 0")
                return
            if shop_name == '':
                self.uiReturn.statusbar.showMessage("Выберите магазин")
                return
            if (float(quantity)+float(rets_quantity)) > float(max_quantity):
                self.uiReturn.statusbar.showMessage("Можно вернуть не более "+str(float(max_quantity) - float(rets_quantity)))
                return
        except:
            self.uiReturn.statusbar.showMessage("Ошибка во введенных данных!")
            return
        connection, cursor = self.connectDB()
        # проверка существования магазина
        cursor.execute("SELECT * FROM shops WHERE number_shops = '" + shop_name + "'")
        names = cursor.fetchall()
        shop_id = -1
        if len(names) == 1:
            shop_id = names[0][0]
        # проверка существования материала
        cursor.execute("SELECT * FROM material WHERE name_material = '" + material + "'")
        names = cursor.fetchall()
        material_id = -1
        if len(names) == 1:
            material_id = names[0][0]
        cursor.execute("SELECT price, quantity  FROM operations_shops WHERE batch = '" + batch + "' AND material = "+str(material_id)+" AND number_shops = "+str(shop_id)+" AND operation = 1")
        names = cursor.fetchall()
        price = -1
        if len(names) == 1:
            price = float(quantity)*(float(names[0][0])/float(names[0][1]))
        print(price)
        connection.close()
        if material_id < 0:
                self.uiReturn.statusbar.showMessage(
                    "Материал не найден!")
                return
        if shop_id < 0:
                self.uiReturn.statusbar.showMessage(
                    "Магазин не найден!")
                return
        if price < 0:
                self.uiReturn.statusbar.showMessage(
                    "Целостность БД нарушена!")
                return
        # занесение сведений в БД
        connection, cursor = self.connectDB()
        cursor.execute(
                "INSERT INTO operations_shops (number_shops, material, quantity, price, batch, operation) VALUES (" +
                str(shop_id) +
                "," + str(material_id) +
                "," + quantity +
                "," + str(price) +
                ",'" + batch +
                "',-1)")
        connection.commit()
        connection.close()
        application.ret.hide()
        application.updateState()
        application.show()

    def updateBatch(self, shop = ''):
        # Table setting
        self.uiReturn.batchT.clear()
        self.uiReturn.batchT.setRowCount(1)
        self.uiReturn.batchT.setColumnCount(4)
        self.uiReturn.batchT.setItem(0, 0, QtWidgets.QTableWidgetItem("№ партии"))
        self.uiReturn.batchT.setItem(0, 1, QtWidgets.QTableWidgetItem("Материал"))
        self.uiReturn.batchT.setItem(0, 2, QtWidgets.QTableWidgetItem("Количество"))
        self.uiReturn.batchT.setItem(0, 3, QtWidgets.QTableWidgetItem("Возвраты"))
        if shop == '': return

        # Take list of batces
        connection, cursor = self.connectDB()
        cursor.execute("SELECT id FROM shops WHERE number_shops = '" + shop +"'")
        names = cursor.fetchall()
        connection.close()
        if len(names) == 0:
            self.uiReturn.statusbar.showMessage("Магазин не существует!")
            return
        shop_id = names[0][0]
        connection, cursor = self.connectDB()
        cursor.execute("""
                SELECT operations_shops.batch, material.name_material, operations_shops.quantity 
                FROM material JOIN operations_shops ON material.id = operations_shops.material 
                WHERE operation = 1 
                AND number_shops = '""" + str(shop_id) +
                """' GROUP BY material.name_material
                """)
        # Add information about shipment and returns
        names = cursor.fetchall()

        self.uiReturn.batchT.setRowCount(len(names) + 1)
        for row, name in enumerate(names):
            self.uiReturn.batchT.setItem(row + 1, 0, QtWidgets.QTableWidgetItem(str(name[0])))
            self.uiReturn.batchT.setItem(row + 1, 1, QtWidgets.QTableWidgetItem(str(name[1])))
            self.uiReturn.batchT.setItem(row + 1, 2, QtWidgets.QTableWidgetItem(str(name[2])))
            cursor.execute("""
                            SELECT operations_shops.quantity 
                            FROM material JOIN operations_shops ON material.id = operations_shops.material 
                            WHERE operation = -1 
                            AND number_shops = '""" + str(shop_id)
                       + "' AND operations_shops.batch = '"+name[0]
                       + "' AND material.name_material = '"+name[1]+"'"
                           )
            # Add information about shipment and returns
            returns = cursor.fetchall()
            ret_count = 0
            for ret in returns:
                ret_count += ret[0]
            self.uiReturn.batchT.setItem(row + 1, 3, QtWidgets.QTableWidgetItem(str(ret_count)))
        connection.close()
        self.uiReturn.batchT.resizeColumnsToContents()

    def updateShop(self):
        connection, cursor = self.connectDB()
        cursor.execute("SELECT number_shops FROM shops")
        names = cursor.fetchall()
        connection.close()
        self.uiReturn.shopsCB.clear()
        for name in names:
            self.uiReturn.shopsCB.addItem(str(name[0]))

if __name__ == '__main__':

    app = QtWidgets.QApplication([])
    application = window()
    application.show()
    sys.exit(app.exec_())

